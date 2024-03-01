import docx
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta

def convert_to_time(updated, pattern = '%d.%m.%Y %H:%M:%S'):
    return datetime.strptime(updated, pattern)

def load_file(filename):
    with open(filename, "r", encoding="utf-8") as f:
        text = f.read()
    return text


def remove_symbols(text):
    # Заменяем все ненужные символы на пустую строку
    return ''.join(c for c in text if c.isprintable())
# Тут функции для обработки данных в XML и формирования docx-файла

#функция конвертирования времени
def convert_time(time_str):
    time_str = time_str.replace(" +0300", "")
    # Разбиваем строку на дату и время
    date_str, time_str = time_str[:8], time_str[8:]
    
    # Добавляем смещение часового пояса
    #time_str = datetime.strptime(time_str, "%H%M%S") + timedelta(hours=3)
    time_str = datetime.strptime(time_str, "%H%M%S")
    
    # Возвращаем время в нужном формате
    time_str = time_str.strftime("%H.%M")
    if time_str[0] == "0":
        time_str = time_str[1:]
    return time_str

# Извлечь дату
def convert_date(time_str):
    date = time_str[:8]
    date = date[6:] + "." + date[4:6] + "." + date[:4]
    return(date)


#функция замены кавычек или обертывания в них
def quotes_change(text):
    if '"' in text:
        # Найти позицию первого вхождения
        index = text.find('"')
        # Заменить первое вхождение
        text = text[:index] + "«" + text[index + 1:]
        # Найти позицию второго вхождения
        index = text.find('"')
        text = text[:index] + "»" + text[index + 1:]    
    else:
        text = chr(171) + text + chr(187)
    return(text)

def prog_create(xml_name, channelID):

    # НАСТРОЙКИ КАНАЛОВ
    # время начала нового дня, например, 450 - это 4:50 утра +- 10 минут
    nastroiki = {'1TV' : {'time': 500, 'unite': ['Новости'], 'zaprety': ["Перерыв в вещании"], 'sub_title': ["ПОДКАСТ.ЛАБ"], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'RTR' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'TVC_RT' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'NTVmsk' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'Piter5_RUS' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'KUL' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'STS' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'TNT' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'TV3' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'RenTV' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'ZVEZDA' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'OTR' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'DOMASHNIY' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'Friday' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    '360d' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'Karusel' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'MIR' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}, 
    'MatchTV' : {'time': 450, 'unite': [], 'zaprety': [], 'sub_title': [], 'hudfilm': 'Фильм', 'serial': 'Сериал'}}



    # обозначаем некоторые важные переменные
    prog_file = "" # В этой переменной будем формировать выходной текст
    previous_title = "" # В этой переменной будет предыдущее название передачи чтобы сразу схлопывать повторы
    programms_in_day = [] # в этом списке будем хранить названия передач, выходящих в течение дня (чтобы не выводить описания у сериалов или фильмов, для которых уже было выведено описание)


    doc = docx.Document()

    # Загружаем XML-файл
    # tree = ET.parse(xml_name)


    # # Получаем корневой элемент
    # root = tree.getroot()
    root = ET.fromstring(xml_name)
    # Находим все элементы "programme"
    programmes = root.findall("programme")

    # Перебираем все элементы "programme"
    for programme in programmes:
        # Получаем информацию о трансляции
        start_time = programme.attrib["start"]
        #stop_time = programme.attrib["stop"]
        #channel = programme.attrib["channel"]
        try:    
            year = programme.find("year").text
        except:
            year = ""
        try:
            country = programme.find("country").text
        except:
            country = ""
        # ищем элемент "rating"
        try:
            rating = programme.find("rating")
            # Получите значение "value"
            value = rating.find("value").text
            value = " (" + value + ")"
        except:
            value = ""
        title = programme.find("title").text
        try:
            ish_sub_title = programme.find("sub-title").text
        except:
            ish_sub_title = ""
        try:
            description = programme.find("desc").text
        except:
            description = ""
        category_elements = programme.findall("category")
        categories = [category.text for category in category_elements]

        # Рассчитываем, не наступил ли новый день начало нового дня
        # если время начала передачи отличается на +-10 минут от константы time_new_day, то выводим надпись о новом дне
        if abs(int(convert_time(start_time).replace(".", "")) - nastroiki[channelID]['time'])<10: 
            doc.add_paragraph()
            doc.add_paragraph(convert_date(start_time) + "\n------------------").bold = True
            doc.add_paragraph()
            previous_title = "" # если наступил новый день, то на всякий случай обнуляем память о предыдущем названии передачи
            programms_in_day = [] # если наступил новый день, то обнуляем список передач за день


        #проверяем, не совпадает ли название новой передачи с предыдущей передачей. Если совпадает - "схлопываем", т.е. не выводим.
        zapret = 0 # регистр, указывающий на присутствие названия программы в списке запретов
        for i in nastroiki[channelID]['zaprety']:
            if i in title:
                zapret = 1
        if title != previous_title and zapret == 0:
            # далее идёт проверка на то, требуется ли для данной конкретной передачи выводить sub-title
            sub_title = ""
            if ish_sub_title:
                for i in nastroiki[channelID]['sub_title']:
                    if i in title:
                        sub_title = ". " + ish_sub_title
            prog_type = "" # переменная для типа передачи (т.к. для худфильмов и сериалов делается особый формат тайтла)
            country_year = "" # переменная для года и страны
            previous_title = title # сохраняем название передачи в переменной "предыдущая передача"
            # Печатаем информацию о трансляции
            #run = doc.add_paragraph().add_run(convert_time(start_time))     
            #run.italic = True
            p = doc.add_paragraph()
            p.add_run(convert_time(start_time)).italic = True   
            p.add_run(chr(9)) # добавили знак табуляции после времени
            for category in categories:
                #print(category) #выводим категории
                # Ищем сериалы или худ/фильмы
                if nastroiki[channelID]['serial'] in category:
                    prog_type = "Сериал"
                if nastroiki[channelID]['hudfilm'] in category:
                    prog_type = "Худфильм"                    
            if prog_type == "Сериал":
                country_year = " (" + country + ", " + year + ")"
                title = quotes_change(title)
                p.add_run(title.upper()).bold = True
            if prog_type == "Худфильм":
                country_year = " (" + country + ", " + year + ")"
                title = quotes_change(title)
                run = p.add_run(title.upper())
                run.bold = True       
                run.italic = True
            if prog_type == "":
                if title.isupper():
                    title = title.capitalize() 
                title = quotes_change(title)
                p.add_run(title)
            p.add_run(sub_title)
            p.add_run(country_year)
            p.add_run(value)
            # Выводим описание для х/ф и сериалов
            if prog_type == "Сериал" and title not in programms_in_day:
                p = doc.add_paragraph()
                p.add_run(description)
            
            programms_in_day.append(title) # добавляем название передачи в список передач за этот день (чтоб потом проверять - нужно ли выводить описание)        

    doc.save(channelID + '_prog.docx')
    return(channelID + '_prog.docx')