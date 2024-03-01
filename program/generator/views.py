# <File>
# <Name>http://xmltv.s-tv.ru/xmltv.php?prg=1708535521&sh=0</Name>
# <EfirWeek>2024-03-04</EfirWeek>
# <Channel>РџСЏС‚РЅРёС†Р°</Channel>
# <ChannelID>Friday</ChannelID>
# <Variant>R</Variant>
# <DateTime>28.02.2024 13:58:43</DateTime>
# </File>
import os
import xml.etree.ElementTree as ET
import requests
import docx

from django.conf import settings
from django.views.generic import FormView, ListView
from django.http import HttpResponseRedirect
from django.urls import reverse

from .forms import GeneratorForm, PROGRAMMES
from .utils import prog_create, remove_symbols
from .models import Record


xml_root = os.path.join(
        settings.MEDIA_ROOT, 'xml'
)

class ChooseProgramFormView(FormView, ):
    form_class = GeneratorForm
    template_name = 'index.html'
    success_url = 'downloads'
    # def get_context_data(self):
    #     d = {}
    #     week = self.request.GET.get('week')
    #     if self.request.GET.get('week'):
    #         d['week'] = week
    #     if self.request.GET.get('programme'):
    #         d['programme'] = week
    #     if self.request.GET.get('variant'):
    #         d['variant'] = week
    #     return d

    def form_valid(self, form):
        """
        If the form is valid, redirect to the supplied URL.
        """
        week = form.cleaned_data.get('week')
        programme = form.cleaned_data.get('programme')
        variant = form.cleaned_data.get('variant')
        if week and not (programme and variant):
            kwargs={'week': week}
        elif week and programme and not variant:
            kwargs={'week': week,'programme':programme }
        elif week and programme and variant:
            kwargs={'week': week,'programme': programme,'variant': variant}
        print(reverse("downloads", args=(week, programme, variant)))
        return HttpResponseRedirect(reverse("downloads", args=(week, programme, variant)))


class ProgramListView(ListView):
    model = Record
    template_name = 'downloads.html'



class ProgramFormView(FormView):
    form_class = GeneratorForm
    template_name = 'index.html'
    success_url = 'downloads'

    def _get_filename(self, week, programme):
        file_name = os.path.join(xml_root, 'xchenel.xml')
        return file_name

    def _get_programme_url_and_channelid(self, filename, week, programme):
        with open(filename, "r", encoding="utf-8") as f:
            text = f.read()
            root = ET.fromstring(text)
            files = root.findall("File")

            for file in files:
                EfirWeek = file.find('EfirWeek').text
                print(EfirWeek)
                ChannelID = file.find('ChannelID').text
                if EfirWeek == week and ChannelID == programme:
                    return file.find('Name').text, file.find('ChannelID').text
        return None, None

    def _get_programme_text(self, programme_url):
        sess = requests.Session()
        payloads = {
            'login':'tv9280',
            'pass': 'TrhB4vlGBC',
        }
        r = sess.get(programme_url, params = payloads)
        xml = r.content.decode('windows-1251')
        return xml
    def post(self, request, *args, **kwargs):
        """
        Handles POST requests, instantiating a form instance with the passed
        POST variables and then checked for validity.
        """
        form_class = self.get_form_class()
        form = self.get_form(form_class)
        if form.is_valid():
            return self.form_valid(form)
        else:
            return self.form_invalid(form)

    def form_valid(self, form):
        """
        If the form is valid, redirect to the supplied URL.
        """
        week = form.cleaned_data['week']
        programme = form.cleaned_data['programme']
        print (f'{week} {programme}')
        filename = self._get_filename(week, programme)
        print(f'file_name={filename}')
        programme_url, channel_id = self._get_programme_url_and_channelid(filename, week, programme)
        print(f'!!!!{programme_url}, {channel_id}')
        if programme_url and channel_id:
            programme_xml = self._get_programme_text(programme_url)
            programme_xml = remove_symbols(programme_xml)
            print(programme_xml[:1000])
            prog_create(programme_xml, channel_id)
        return HttpResponseRedirect(self.get_success_url())

def download_file():
    doc = docx.Document()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_run('description')
    doc.save()
    pass